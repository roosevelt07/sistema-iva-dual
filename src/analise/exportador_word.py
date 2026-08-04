"""Exportação do relatório de análise em Word (.docx)."""

import io
from decimal import Decimal
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

from src.analise.decisao import ResultadoAnalise
from src.analise.formatters import fmt_moeda
from src.analise.narrativa import gerar_narrativa

LOGO_PATH = Path(__file__).parent.parent.parent / "assets" / "logo.png"

_VERDE = RGBColor(0x00, 0xB2, 0xA9) if DOCX_DISPONIVEL else None
_VERDE_HEX = "00B2A9"

# Labels em português para as chaves internas de `cenario.detalhes`,
# usadas na Memória de cálculo. Cobre as chaves reais dos 3 regimes
# (src/analise/regimes/{simples,presumido,real}.py) + cenário B
# compartilhado (_cenario_b_regime_regular). carga_bruta/carga_liquida/
# credito_aproveitado são campos de CenarioUnico, não chaves de
# `.detalhes` — mantidos aqui só por completude, nunca aparecem nesta
# tabela hoje.
LABELS_PT: dict = {
    "ibs_cbs_das":            "IBS/CBS no DAS (Simples)",
    "credito_entrada":        "Crédito de entrada",
    "ibs_cbs_bruto":          "IBS/CBS bruto (regime regular)",
    "credito_regime_normal":  "Crédito — fornecedores regime normal",
    "credito_simples":        "Crédito — fornecedores Simples",
    "pis_cofins":             "PIS/COFINS",
    "pis_cofins_bruto":       "PIS/COFINS",
    "credito_pis_cofins":     "Crédito PIS/COFINS",
    "icms":                   "ICMS",
    "iss":                    "ISS",
    "reducao_icms_iss_aplicada": "Redução ICMS/ISS aplicada",
    "carga_bruta":            "Carga bruta",
    "carga_liquida":          "Carga líquida",
    "credito_aproveitado":    "Crédito aproveitado",
}


def _detalhes_monetarios(detalhes: dict) -> dict:
    return {k: v for k, v in detalhes.items() if isinstance(v, Decimal)}


def _titulo_verde(doc, texto, level):
    heading = doc.add_heading(texto, level=level)
    heading.runs[0].font.color.rgb = _VERDE
    return heading


def _adicionar_borda_paragrafo(paragrafo, lados=("top", "left", "bottom", "right")) -> None:
    """Borda teal ao redor (ou só em alguns lados) do parágrafo."""
    p_pr = paragrafo._p.get_or_add_pPr()
    borda = OxmlElement("w:pBdr")
    for lado in lados:
        elemento = OxmlElement(f"w:{lado}")
        elemento.set(qn("w:val"), "single")
        elemento.set(qn("w:sz"), "12")
        elemento.set(qn("w:color"), _VERDE_HEX)
        elemento.set(qn("w:space"), "4")
        borda.append(elemento)
    p_pr.append(borda)


def _configurar_cabecalho(doc, nome_cliente: str) -> None:
    header = doc.sections[0].header
    p_logo = header.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        p_logo.add_run().add_picture(str(LOGO_PATH), height=Inches(0.5))

    p_linha = header.add_paragraph()
    _adicionar_borda_paragrafo(p_linha, lados=("bottom",))


def _configurar_rodape(doc, resultado: ResultadoAnalise) -> None:
    rodape = doc.sections[0].footer.paragraphs[0]
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    extrato = getattr(resultado, "extrato", None)
    cnpj = getattr(extrato, "cnpj_basico", None)
    pa = getattr(extrato, "periodo_apuracao", None)
    if cnpj or pa:
        partes = []
        if cnpj:
            partes.append(f"CNPJ: {cnpj}")
        if pa:
            partes.append(f"PA: {pa}")
        rodape.add_run(" | ".join(partes) + "  —  ").italic = True

    rodape.add_run(
        "Simulação baseada na LC 214/2025 — valores estimados."
    ).italic = True


def gerar_word(resultado: ResultadoAnalise, nome_cliente: str, data_relatorio, observacoes: str = "") -> bytes:
    """Gera o relatório Word e retorna os bytes do arquivo .docx."""
    if not DOCX_DISPONIVEL:
        raise RuntimeError("python-docx não está instalado.")

    cenario_a = resultado.cenarios.cenario_a
    cenario_b = resultado.cenarios.cenario_b

    doc = Document()
    _configurar_cabecalho(doc, nome_cliente)
    _configurar_rodape(doc, resultado)

    # 1. Capa
    if LOGO_PATH.exists():
        p_logo_capa = doc.add_paragraph()
        p_logo_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo_capa.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
    titulo = doc.add_heading("Análise IVA Dual — LC 214/2025", level=0)
    titulo.runs[0].font.color.rgb = _VERDE
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Cliente: {nome_cliente or '—'}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    extrato = getattr(resultado, "extrato", None)
    cnpj = getattr(extrato, "cnpj_basico", None)
    if cnpj:
        p = doc.add_paragraph(f"CNPJ: {cnpj}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Data: {data_relatorio}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Resumo executivo
    _titulo_verde(doc, "Resumo executivo", level=1)
    doc.add_paragraph(gerar_narrativa(resultado))

    # 3. Comparativo de carga
    _titulo_verde(doc, "Comparativo de carga tributária", level=1)
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Light Grid Accent 1"
    for i, texto in enumerate(["Item", "Atual (R$)", "IVA (R$)", "Economia (R$)"]):
        tabela.rows[0].cells[i].text = texto
    linhas = [
        ("Carga bruta", cenario_a.carga_bruta, cenario_b.carga_bruta),
        ("Crédito aproveitado", cenario_a.credito_aproveitado, cenario_b.credito_aproveitado),
        ("Carga líquida", cenario_a.carga_liquida, cenario_b.carga_liquida),
    ]
    for nome, valor_a, valor_b in linhas:
        linha = tabela.add_row().cells
        linha[0].text = nome
        linha[1].text = fmt_moeda(valor_a)
        linha[2].text = fmt_moeda(valor_b)
        linha[3].text = fmt_moeda(valor_a - valor_b)

    # 4. Atividades declaradas
    _titulo_verde(doc, "Atividades declaradas", level=1)
    blocos_atividade = getattr(resultado.ctx, "blocos_atividade", None)
    if blocos_atividade:
        tabela_ativ = doc.add_table(rows=1, cols=5)
        tabela_ativ.style = "Light Grid Accent 1"
        for i, texto in enumerate(["Atividade", "Anexo", "Receita", "DAS", "Obs."]):
            tabela_ativ.rows[0].cells[i].text = texto
        for bloco in blocos_atividade:
            linha = tabela_ativ.add_row().cells
            linha[0].text = getattr(bloco, "nome", "")
            linha[1].text = str(getattr(bloco, "anexo", ""))
            linha[2].text = fmt_moeda(getattr(bloco, "receita_atividade", Decimal("0")))
            linha[3].text = fmt_moeda(getattr(bloco, "das", Decimal("0")))
            linha[4].text = getattr(bloco, "observacao", "")
        anexo_predominante = getattr(resultado, "anexo_predominante", None)
        if anexo_predominante:
            doc.add_paragraph(f"Anexo predominante: {anexo_predominante}")
    else:
        doc.add_paragraph("Dados disponíveis no Modo 3 — Extrato PGDAS-D.")

    # 5. Recomendação
    _titulo_verde(doc, "Recomendação", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "Migrar para o regime regular do IBS/CBS."
        if resultado.recomendacao == "MIGRAR_REGIME_REGULAR"
        else "Manter o regime atual."
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _VERDE
    _adicionar_borda_paragrafo(p)

    # 6. Alertas e ressalvas
    if resultado.alertas:
        _titulo_verde(doc, "Alertas e ressalvas", level=1)
        for alerta in resultado.alertas:
            doc.add_paragraph(alerta, style="List Bullet")

    if observacoes:
        _titulo_verde(doc, "Observações adicionais", level=1)
        doc.add_paragraph(observacoes)

    # 7. Memória de cálculo
    _titulo_verde(doc, "Memória de cálculo", level=1)
    for cenario in (cenario_a, cenario_b):
        doc.add_heading(cenario.nome, level=2)
        detalhes = _detalhes_monetarios(cenario.detalhes)
        tabela_mem = doc.add_table(rows=1, cols=2)
        tabela_mem.style = "Light List Accent 1"
        tabela_mem.rows[0].cells[0].text = "Item"
        tabela_mem.rows[0].cells[1].text = "Valor"
        for chave, valor in detalhes.items():
            label = LABELS_PT.get(chave, chave.replace("_", " ").title())
            linha = tabela_mem.add_row().cells
            linha[0].text = label
            linha[1].text = fmt_moeda(valor)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
